
import sys
import os
import datetime

import pandas as pd
import re
import docx
from collections import Counter
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QApplication, QLabel, QMainWindow, QMenuBar, QMenu,\
    QToolBar, QAction, QComboBox, QPushButton, QTableWidget, QFileDialog, QTableWidgetItem, QInputDialog, QDialog,\
    QDialogButtonBox,QMessageBox, QWidget, QGridLayout
from PyQt5.QtGui import QIcon
from docx.shared import Pt
class Window(QMainWindow):
    """Main Window."""
    def __init__(self, parent=None):
        """Initializer."""
        super().__init__(parent)
        self.setWindowTitle("DNA/RNA Calc")
        self.setFixedSize(650, 450)
        self.lbl = QLabel("Путь к файлу: [файл не выбран]",self)
        self.lbl.move(20,30)
        self.lbl.adjustSize()
        self.setWindowIcon(QIcon("icon2.png"))
        self.tempFileLoc = ['']
        self._creatTableWidgetDry()
        self._creatTableWidgetLiq()
        self._creatAction()
        self._creatMenuBar()
        self._connectAction()
        self.RowCount = [0]
        self.RowCountLiq = [0]
        self._creatComboBox()
        self._creatPushButtom2()
        self._creatPushButtom3()
        self._creatPushButtom4()
        self._creatPushButtom5()
        self._creatPushButtom6()
        self._creatPushButtom7()
    def _creatMenuBar(self):
        menuBar = self.menuBar()
        filemenu = QMenu("File", self)
        menuBar.addMenu(filemenu)
        filemenu.addAction(self.newAction)
        Sample = filemenu.addMenu("Create Sample")
        Sample.addAction(self.sampleDry)
        Sample.addAction(self.sampleLiq)
        filemenu.addAction(self.openAction)
        filemenu.addAction(self.saveAction)
        saveAs = filemenu.addMenu("Save as")
        saveAs.setIcon(QIcon("save-file.png"))
        saveAs.addAction(self.saveAsActionOnlyExel)
        saveAs.addAction(self.saveAsActionWithWord)
        filemenu.addAction(self.exitAction)
    def _creatAction(self):
        self.newAction = QAction(self)
        self.newAction.setText("New")
        self.sampleDry = QAction(self)
        self.sampleDry.setText("Dry")
        self.sampleLiq = QAction(self)
        self.sampleLiq.setText("Liquid")
        self.openAction = QAction(self)
        self.openAction.setText("Open")
        self.openAction.setIcon(QIcon("iconOpenFile.png"))
        self.saveAction = QAction(self)
        self.saveAction.setText("Save")
        self.saveAction.setIcon(QIcon("iconSaveFile.png"))
        self.saveAsActionOnlyExel = QAction(self)
        self.saveAsActionOnlyExel.setText("Only Exel")
        self.saveAsActionWithWord = QAction(self)
        self.saveAsActionWithWord.setText("Word and Exel")
        self.exitAction = QAction(self)
        self.exitAction.setText("Exit")
        self.exitAction.setIcon(QIcon("logout.png"))
    def newFile(self):
        self.tempFileLoc = ['']
        self.lbl.setText("Путь к файлу: [файл не выбран]")
    def creatSampleDry(self):
        fname = QFileDialog.getSaveFileName(self,"Save file","","Xlsx Files (*.xlsx)")
        if fname != " ":
            df = pd.DataFrame(columns=['Назание', 'последовательность', 'Длина', 'Мол. масса', 'OE', 'Пмоль', 'Мкг'])
            df2 = pd.DataFrame(columns=['Название','Зонд', 'Последовательность','Зонд', 'OE', 'DNA/RNA'])
            with pd.ExcelWriter(fname[0], engine='openpyxl') as writer:
                df.to_excel(writer,sheet_name='Sheet1')
                df2.to_excel(writer,sheet_name= 'Sheet2')
    def creatSampleLiquid(self):
        fname = QFileDialog.getSaveFileName(self, "Save file", "", "Xlsx Files (*.xlsx)")
        if fname != " ":
            df = pd.DataFrame(columns=['Назание', 'последовательность', 'Длина', 'Мол. масса', 'OE', 'мкл', 'OE/ml','Пмоль/mkl', 'Мкг/mkl'])
            df2 = pd.DataFrame(columns=['Название', 'Зонд', 'Последовательность', 'Зонд', 'OE','pmole/mkl', 'DNA/RNA'])
            with pd.ExcelWriter(fname[0], engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Sheet1')
                df2.to_excel(writer, sheet_name='Sheet2')
    def openFile(self):
        fname = QFileDialog()
        tempFileLoc = fname.getOpenFileName(self, 'Open file', ' ', 'Xlsx Files (*.xlsx)')
        if tempFileLoc[1] == '':
            pass
        else:
            try:
                self.tempFileLoc = tempFileLoc
                self.pathToFile(tempFileLoc)
                tempdf = pd.DataFrame(pd.read_excel(self.tempFileLoc[0],sheet_name=1,engine='openpyxl'))
                if len(tempdf.axes[1])-1 <= self.table.columnCount():
                    tempdf = pd.DataFrame(pd.read_excel(self.tempFileLoc[0],sheet_name=1,engine='openpyxl',usecols="B:G"))
                    df = tempdf.astype(str)
                    if self.table.isVisible() == False:
                        self._feature4()
                    for i in range(0, len(df.axes[0])):
                        self._feature1()
                        if df.iat[i,5] == "RNA":
                            self.comboBoxCount[i].setCurrentIndex(1)
                        for j in range(0, len(df.axes[1])):
                            if pd.isna(df.iat[i,j]) == False:
                                self.table.setItem(i,j,QTableWidgetItem(df.iat[i,j]))
                elif len(tempdf.axes[1])-1 <= self.tableLiq.columnCount():
                    tempdf = pd.DataFrame(pd.read_excel(self.tempFileLoc[0], sheet_name=1, engine='openpyxl', usecols="B:H"))
                    df = tempdf.astype(str)
                    if self.tableLiq.isVisible() == False:
                        self._feature3()
                    for i in range(0, len(df.axes[0])):
                        self._feature1()
                        if df.iat[i,6] == "RNA":
                            self.comboBoxCountLiq[i].setCurrentIndex(1)
                        for j in range(0, len(df.axes[1])):
                            if pd.isna(df.iat[i,j]) == False:
                                self.tableLiq.setItem(i,j,QTableWidgetItem(df.iat[i,j]))
            except:
                self._creatWarningWindow("Выбран некорректный файл!")
    def saveFile(self,filename,word_or_not):
        try:
            temp3 = []
            if self.table.isVisible() == True:
                df = pd.DataFrame(columns=['Назание','последовательность','Длина','Мол. масса','OE','Пмоль','Мкг'])
                df2 = pd.DataFrame(columns=['Название','Зонд','Последовательность','Зонд','OE','DNA/RNA'])
                for j in range(0,self.RowCount[-1]):
                    if self.table.item(j,2) != None:
                        temp3.append(re.sub(r'\s*\{[^{}]*\}',"",self.table.item(j,2).text()))
                    else:
                        temp3.append(self.table.item(j, 2))
                for i in range(0,len(temp3)):
                    temp5 = Counter(temp3[i])
                    if self.comboBoxCount[i].currentText() == "DNA":
                        temp6 = (float(self.table.item(i,4).text().replace(',','.'))*(100/(((temp5['A']+temp5['a']) * 1.53)
                                                                                           + ((temp5['R'] + temp5['r']) * 1.35)
                                                                                           + ((temp5['Y'] + temp5['y']) * 0.775)
                                                                                           + ((temp5['M'] + temp5['m']) * 1.105)
                                                                                           + ((temp5['K'] + temp5['k']) * 1.02)
                                                                                           + ((temp5['G'] + temp5['g'] + temp5['I'] + temp5['i']) * 1.18)
                                                                                           + ((temp5['S'] + temp5['s']) * 0.955)
                                                                                           + ((temp5['W'] + temp5['w']) * 1.17)
                                                                                           + ((temp5['B'] + temp5['b']) * 0.917)
                                                                                           + ((temp5['C'] + temp5['c']) * 0.74)
                                                                                           + ((temp5['D'] + temp5['d']) * 1.18)
                                                                                           + ((temp5['U'] + temp5['u']) * 1.017)
                                                                                           + ((temp5['V'] + temp5['v']) * 1.137)
                                                                                           + ((temp5['N'] + temp5['n']) * 1.137)
                                                                                           + ((temp5['T'] + temp5['t']) * 0.93)))*1000)
                        temp7 = (((temp5['A']+temp5['a']) * 313.21)
                                 + ((temp5['T'] + temp5['t']) * 304.2)
                                 + ((temp5['C'] + temp5['c']) * 289.18)
                                 + ((temp5['G'] + temp5['g'] + temp5['I'] + temp5['i']) * 329.21)
                                 + ((temp5['Y'] + temp5['y']) * 296.66)
                                 + ((temp5['M'] + temp5['m']) * 301.16)
                                 + ((temp5['R'] + temp5['r']) * 321.21)
                                 + ((temp5['K'] + temp5['k'] + temp5['D'] + temp5['d']) * 316.71)
                                 + ((temp5['U'] + temp5['u']) * 290.16)
                                 + ((temp5['S'] + temp5['s'] + temp5['V'] + temp5['v'] + temp5['N'] + temp5['n']) * 309.21)
                                 + ((temp5['W'] + temp5['w']) * 308.66)
                                 + ((temp5['B'] + temp5['b']) * 309.15) - 61.96)
                        if self.table.item(i,1) != None and self.table.item(i,3) != None:
                            if (self.table.item(i,1).text() != "" and self.table.item(i,3).text() != "") and (
                                    self.table.item(i,1).text() != "nan" and self.table.item(i,3).text() != "nan"):
                                df.loc[i] = [self.table.item(i,0).text(),
                                            self.table.item(i,1).text() + "-" + re.sub(r'[{}]', "", self.table.item(i, 2).text()) + "-" + self.table.item(i,3).text(),
                                            sum(map(str.isalpha, self.table.item(i,2).text())),
                                            round(temp7),
                                            self.table.item(i,4).text(),
                                            int(round(temp6,-1)),
                                            str(float(round((temp6*temp7*0.000001), 1))).replace('.',',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                            self.table.item(i,1).text(),
                                            self.table.item(i, 2).text(),
                                            self.table.item(i,3).text(),
                                            self.table.item(i,4).text(),
                                            self.comboBoxCount[i].currentText()]
                            elif (self.table.item(i,1).text() != "" and self.table.item(i,3).text() == "") or (
                                    self.table.item(i,1).text() != "nan" and self.table.item(i,3).text() == "nan"):
                                    df.loc[i] = [self.table.item(i, 0).text(),
                                                self.table.item(i, 1).text() + "-" + re.sub(r'[{}]', "", self.table.item(i, 2).text()),
                                                sum(map(str.isalpha, self.table.item(i, 2).text())),
                                                round(temp7),
                                                self.table.item(i, 4).text(),
                                                int(round(temp6, -1)),
                                                str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                    df2.loc[i] = [self.table.item(i, 0).text(),
                                                self.table.item(i, 1).text(),
                                                self.table.item(i, 2).text(),
                                                "",
                                                self.table.item(i, 4).text(),
                                                self.comboBoxCount[i].currentText()]
                            elif (self.table.item(i,1).text() == "" and self.table.item(i,3).text() != "") or (
                                    self.table.item(i,1).text() == "nan" and self.table.item(i,3).text() != "nan"):
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             re.sub(r'[{}]', "", self.table.item(i, 2).text()) + "-" + self.table.item(i, 3).text(),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              "",
                                              self.table.item(i, 2).text(),
                                              self.table.item(i, 3).text(),
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]
                            elif (self.table.item(i,1).text() == "" and self.table.item(i,3).text() == "") or (
                                    self.table.item(i,1).text() == "nan" and self.table.item(i,3).text() == "nan"):
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             re.sub(r'[{}]', "", self.table.item(i, 2).text()),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              "",
                                              self.table.item(i, 2).text(),
                                              "",
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]

                        elif self.table.item(i,1) == None and self.table.item(i,3) != None:
                            if self.table.item(i,3).text() != "" and self.table.item(i,3).text() != "nan":
                                df.loc[i] = [self.table.item(i,0).text(),
                                            re.sub(r'[{}]', "", self.table.item(i, 2).text()) + "-" + self.table.item(i,3).text(),
                                            sum(map(str.isalpha, self.table.item(i,2).text())),
                                            round(temp7),
                                            self.table.item(i,4).text(),
                                            int(round(temp6,-1)),
                                            str(float(round((temp6*temp7*0.000001), 1))).replace('.',',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                            "",
                                            self.table.item(i, 2).text(),
                                            self.table.item(i,3).text(),
                                            self.table.item(i,4).text(),
                                            self.comboBoxCount[i].currentText()]
                        elif self.table.item(i,1) != None and self.table.item(i,3) == None:
                            if self.table.item(i,1).text() != "" and self.table.item(i,1).text() != "nan":
                                df.loc[i] = [self.table.item(i,0).text(),
                                            self.table.item(i,1).text() + "-" + re.sub(r'[{}]', "", self.table.item(i, 2).text()),
                                            sum(map(str.isalpha, self.table.item(i,2).text())),
                                            round(temp7),
                                            self.table.item(i,4).text(),
                                            int(round(temp6,-1)),
                                            str(float(round((temp6*temp7*0.000001), 1))).replace('.',',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                            self.table.item(i,1).text(),
                                            self.table.item(i, 2).text(),
                                            "",
                                            self.table.item(i,4).text(),
                                            self.comboBoxCount[i].currentText()]
                        elif self.table.item(i,1) == None and self.table.item(i,3) == None:
                            df.loc[i] = [self.table.item(i,0).text(),
                                        re.sub(r'[{}]', "", self.table.item(i, 2).text()),
                                        sum(map(str.isalpha, self.table.item(i,2).text())),
                                        round(temp7),
                                        self.table.item(i,4).text(),
                                        int(round(temp6,-1)),
                                        str(float(round((temp6*temp7*0.000001), 1))).replace('.',',')]
                            df2.loc[i] = [self.table.item(i, 0).text(),
                                        "",
                                        self.table.item(i, 2).text(),
                                        "",
                                        self.table.item(i,4).text(),
                                        self.comboBoxCount[i].currentText()]
                    elif self.comboBoxCount[i].currentText() == "RNA":
                        temp6 = (float(self.table.item(i,4).text().replace(',','.'))*(100/(((temp5['A']+temp5['a']) * 1.53) + ((temp5['G']+temp5['g']) * 1.18) + ((temp5['C']+temp5['c']) * 0.74) + ((temp5['T']+temp5['U']+temp5['t']+temp5['u']) * 0.93)))*1000)
                        temp7 = (((temp5['A']+temp5['a'])*329.21)+((temp5['T']+temp5['t'])*320.2)+((temp5['C']+temp5['c'])*305.18)+((temp5['G']+temp5['g'])*345.21)+((temp5['U']+temp5['u'])*306.17)+159.0)
                        if self.table.item(i, 1) != None and self.table.item(i, 3) != None:
                            if (self.table.item(i, 1).text() != "" and self.table.item(i, 3).text() != "") and (
                                    self.table.item(i, 1).text() != "nan" and self.table.item(i, 3).text() != "nan"):
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             self.table.item(i, 1).text() + "-" + re.sub(r'[{}]', "", self.table.item(i,
                                                                                                                      2).text()) + "-" + self.table.item(
                                                 i, 3).text(),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              self.table.item(i, 1).text(),
                                              self.table.item(i, 2).text(),
                                              self.table.item(i, 3).text(),
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]
                            elif (self.table.item(i, 1).text() != "" and self.table.item(i, 3).text() == "") or (
                                    self.table.item(i, 1).text() != "nan" and self.table.item(i, 3).text() == "nan"):
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             self.table.item(i, 1).text() + "-" + re.sub(r'[{}]', "",
                                                                                         self.table.item(i, 2).text()),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              self.table.item(i, 1).text(),
                                              self.table.item(i, 2).text(),
                                              "",
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]
                            elif (self.table.item(i, 1).text() == "" and self.table.item(i, 3).text() != "") or (
                                    self.table.item(i, 1).text() == "nan" and self.table.item(i, 3).text() != "nan"):
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             re.sub(r'[{}]', "", self.table.item(i, 2).text()) + "-" + self.table.item(
                                                 i, 3).text(),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              "",
                                              self.table.item(i, 2).text(),
                                              self.table.item(i, 3).text(),
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]
                            elif (self.table.item(i, 1).text() == "" and self.table.item(i, 3).text() == "") or (
                                    self.table.item(i, 1).text() == "nan" and self.table.item(i, 3).text() == "nan"):
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             re.sub(r'[{}]', "", self.table.item(i, 2).text()),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              "",
                                              self.table.item(i, 2).text(),
                                              "",
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]

                        elif self.table.item(i, 1) == None and self.table.item(i, 3) != None:
                            if self.table.item(i, 3).text() != "" and self.table.item(i, 3).text() != "nan":
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             re.sub(r'[{}]', "", self.table.item(i, 2).text()) + "-" + self.table.item(
                                                 i, 3).text(),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              "",
                                              self.table.item(i, 2).text(),
                                              self.table.item(i, 3).text(),
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]
                        elif self.table.item(i, 1) != None and self.table.item(i, 3) == None:
                            if self.table.item(i, 1).text() != "" and self.table.item(i, 1).text() != "nan":
                                df.loc[i] = [self.table.item(i, 0).text(),
                                             self.table.item(i, 1).text() + "-" + re.sub(r'[{}]', "",
                                                                                         self.table.item(i, 2).text()),
                                             sum(map(str.isalpha, self.table.item(i, 2).text())),
                                             round(temp7),
                                             self.table.item(i, 4).text(),
                                             int(round(temp6, -1)),
                                             str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                                df2.loc[i] = [self.table.item(i, 0).text(),
                                              self.table.item(i, 1).text(),
                                              self.table.item(i, 2).text(),
                                              "",
                                              self.table.item(i, 4).text(),
                                              self.comboBoxCount[i].currentText()]
                        elif self.table.item(i, 1) == None and self.table.item(i, 3) == None:
                            df.loc[i] = [self.table.item(i, 0).text(),
                                         re.sub(r'[{}]', "", self.table.item(i, 2).text()),
                                         sum(map(str.isalpha, self.table.item(i, 2).text())),
                                         round(temp7),
                                         self.table.item(i, 4).text(),
                                         int(round(temp6, -1)),
                                         str(float(round((temp6 * temp7 * 0.000001), 1))).replace('.', ',')]
                            df2.loc[i] = [self.table.item(i, 0).text(),
                                          "",
                                          self.table.item(i, 2).text(),
                                          "",
                                          self.table.item(i, 4).text(),
                                          self.comboBoxCount[i].currentText()]
                with pd.ExcelWriter(filename[0],engine='openpyxl') as writer:
                    df.to_excel(writer,sheet_name='Sheet1')
                    df2.to_excel(writer,sheet_name='Sheet2')
                print(df2)
                if word_or_not == True:
                    doc = docx.Document('Sample.docx')
                    table = doc.tables[0]
                    for i in range(df.shape[0]):
                        table.add_row().cells

                        table.cell(i + 2, 0).text = str(i + 1)
                        table.cell(i + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 1).text = str(df.values[i, 0])
                        table.cell(i + 2, 1).paragraphs[0].runs[0].font.bold = True
                        table.cell(i + 2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if (df2.values[i,1] == "nan" and df2.values[i,3] == "nan") or (df2.values[i,1] == "" and df2.values[i,3] == "") or (df2.values[i,1] == None and df2.values[i,3] == None):
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif (df2.values[i, 1] == "nan" and df2.values[i, 3] != "nan") or (df2.values[i,1] == "" and df2.values[i,3] != "") or (df2.values[i,1] == None and df2.values[i,3] != None):
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i,2]))
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i,3])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif (df2.values[i, 1] != "nan" and df2.values[i, 3] == "nan") or (df2.values[i,1] != "" and df2.values[i,3] == "") or (df2.values[i,1] != None and df2.values[i,3] == None):
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 1])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 1])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i,3])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                        table.cell(i + 2, 3).text = str(df.values[i, 2])
                        table.cell(i + 2, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 4).text = str(df.values[i, 3])
                        table.cell(i + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 5).text = str(df.values[i, 4])
                        table.cell(i + 2, 5).paragraphs[0].runs[0].font.bold = True
                        table.cell(i + 2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 6).text = str(df.values[i, 5])
                        table.cell(i + 2, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 7).text = str(df.values[i, 6])
                        table.cell(i + 2, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 8).text = "ПААГ"
                        table.cell(i + 2, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.style = 'Table Grid'
                    print('doc')
                    f_splited = filename[0].split('/')
                    fname = f_splited[len(f_splited) - 1][:-5]
                    print(fname)
                    doc.save(filename[0] + '.docx')
                    doc = docx.Document('Stickers.docx')
                    table = doc.tables[0]
                    row = 0
                    column = 0
                    temp_row = 1
                    now = datetime.datetime.now()
                    for i in range(df.shape[0]):
                        table.cell(row, column).add_paragraph(
                            str(df.values[i, 0])).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        table.cell(row, column).paragraphs[1].runs[0].font.bold = True
                        table.cell(row, column).paragraphs[1].runs[0].font.name = 'Arial'
                        if (df2.values[i,1] == "nan" and df2.values[i,3] == "nan") or (df2.values[i,1] == "" and df2.values[i,3] == "") or (df2.values[i,1] == None and df2.values[i,3] == None):
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(" 5'-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i,2]))
                            table.cell(row, column).paragraphs[2].add_run("-3' ").font.bold = True
                        elif (df2.values[i, 1] == "nan" and df2.values[i, 3] != "nan") or (df2.values[i,1] == "" and df2.values[i,3] != "") or (df2.values[i,1] == None and df2.values[i,3] != None):
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(" 5'-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run("-" + str(df2.values[i,3]) + " ").font.bold = True
                        elif (df2.values[i, 1] != "nan" and df2.values[i, 3] == "nan") or (df2.values[i,1] != "" and df2.values[i,3] == "") or (df2.values[i,1] != None and df2.values[i,3] == None):
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(str(" " + df2.values[i,1]) + "-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run("-3' ").font.bold = True
                        else:
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(" " + str(df2.values[i, 1]) + "-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run("-" + str(df2.values[i, 3]) + " ").font.bold = True
                        table.cell(row, column).add_paragraph(
                            str(df.values[i, 4] + " OE")).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        table.cell(row, column).paragraphs[3].runs[0].font.bold = True
                        table.cell(row, column).add_paragraph(
                            " Stor.-20°C        " + now.strftime('%d %b %Y ')).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        column += 2
                        if column > 6:
                            row += 1
                            temp_row += 1
                            column = 0
                    for i in range(temp_row):
                        for j in range(0, 7, 2):
                            paragraphs = table.cell(i, j).paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size = Pt(6)
                    doc.save(filename[0] + 'Наклейки.docx')

            elif self.tableLiq.isVisible() == True:
                df = pd.DataFrame(columns=['Назание', 'последовательность', 'Длина', 'Мол. масса', 'OE', 'мкл', 'OE/ml','Пмоль/mkl', 'Мкг/mkl'])
                df2 = pd.DataFrame(columns=['Название','Зонд', 'Последовательность','Зонд', 'OE','пмоль/мкг','DNA/RNA'])
                for j in range(0,self.RowCountLiq[-1]):
                    if self.tableLiq.item(j,2) != None:
                        temp3.append(re.sub(r'\s*\{[^{}]*\}',"",self.tableLiq.item(j,2).text()))
                        print(self.tableLiq.item(0,4).text().replace(',','.'))
                        print(1 + float(self.tableLiq.item(0,4).text().replace(',','.')))
                    else:
                        print('huy1')
                        temp3.append(self.tableLiq.item(j, 2))
                for i in range(0,len(temp3)):
                    print('huy2')
                    temp5 = Counter(temp3[i])
                    if self.comboBoxCountLiq[i].currentText() == "DNA":
                        print('huy3')
                        temp6 = (float(self.tableLiq.item(i,4).text().replace(',','.'))*(100/(((temp5['A']+temp5['a']) * 1.53)
                                                                                              + ((temp5['R'] + temp5['r'])*1.35)
                                                                                              + ((temp5['Y'] + temp5['y'])*0.775)
                                                                                              + ((temp5['M'] + temp5['m'])*1.105)
                                                                                              + ((temp5['K'] + temp5['k'])*1.02)
                                                                                              + ((temp5['G'] + temp5['g'] + temp5['I'] + temp5['i']) * 1.18)
                                                                                              + ((temp5['S'] + temp5['s']) * 0.955)
                                                                                              + ((temp5['W'] + temp5['w']) * 1.17)
                                                                                              + ((temp5['B'] + temp5['b']) * 0.917)
                                                                                              + ((temp5['C'] + temp5['c']) * 0.74)
                                                                                              + ((temp5['D'] + temp5['d']) * 1.18)
                                                                                              + ((temp5['U'] + temp5['u']) * 1.017)
                                                                                              + ((temp5['V'] + temp5['v']) * 1.137)
                                                                                              + ((temp5['N'] + temp5['n']) * 1.137)
                                                                                              + ((temp5['T'] + temp5['t']) * 0.93)))*1000/int(self.tableLiq.item(i,5).text()))
                        temp7 = (((temp5['A'] + temp5['a']) * 313.21)
                                 + ((temp5['T'] + temp5['t']) * 304.2)
                                 + ((temp5['C'] + temp5['c']) * 289.18)
                                 + ((temp5['G'] + temp5['g'] + temp5['I'] + temp5['i']) * 329.21)
                                 + ((temp5['Y'] + temp5['y']) * 296.66)
                                 + ((temp5['M'] + temp5['m']) * 301.16)
                                 + ((temp5['R'] + temp5['r']) * 321.21)
                                 + ((temp5['K'] + temp5['k'] + temp5['D'] + temp5['d']) * 316.71)
                                 + ((temp5['U'] + temp5['u']) * 290.16)
                                 + ((temp5['S'] + temp5['s'] + temp5['V'] + temp5['v'] + temp5['N'] + temp5['n']) * 309.21)
                                 + ((temp5['W'] + temp5['w']) * 308.66)
                                 + ((temp5['B'] + temp5['b']) * 309.15) - 61.96)
                        temp8 = (float(self.tableLiq.item(i,4).text().replace(',','.'))/temp6*1000)
                        temp9 = (float(self.tableLiq.item(i,5).text())*temp7*0.000001)
                        if (self.tableLiq.item(i, 1) != None and self.tableLiq.item(i, 3) != None):
                            if (self.tableLiq.item(i, 1).text() != "" and self.tableLiq.item(i, 3).text() != "") and (
                                     self.tableLiq.item(i, 1).text() != "nan" and self.tableLiq.item(i, 3).text() != "nan"):
                                df.loc[i] = [self.tableLiq.item(i,0).text(),
                                             self.tableLiq.item(i,1).text() + "-" + re.sub(r'[\{\}]', "", self.tableLiq.item(i, 2).text()) + "-" + self.tableLiq.item(i,3).text(),
                                             sum(map(str.isalpha, self.tableLiq.item(i,2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i,4).text(),
                                             str(float(round(temp6,1))).replace('.',','),
                                             str(float(round(temp8, 1))).replace('.',','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.',',')]
                                df2.loc[i] = [self.tableLiq.item(i,0).text(),
                                              self.tableLiq.item(i,1).text(),
                                              self.tableLiq.item(i,2).text(),
                                              self.tableLiq.item(i,3).text(),
                                              self.tableLiq.item(i,4).text(),
                                              self.tableLiq.item(i,5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                            elif (self.tableLiq.item(i, 1).text() == "" and self.tableLiq.item(i, 3).text() != "") or (
                                     self.tableLiq.item(i, 1).text() == "nan" and self.tableLiq.item(i, 3).text() != "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             re.sub(r'[\{\}]', "",self.tableLiq.item(i, 2).text()) + "-" + self.tableLiq.item(i,3).text(),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              "",
                                              self.tableLiq.item(i, 2).text(),
                                              self.tableLiq.item(i, 3).text(),
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                            elif (self.tableLiq.item(i, 1).text() != "" and self.tableLiq.item(i, 3).text() == "") or (
                                     self.tableLiq.item(i, 1).text() != "nan" and self.tableLiq.item(i, 3).text() == "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             self.tableLiq.item(i, 1).text() + "-" + re.sub(r'[\{\}]', "",self.tableLiq.item(i,2).text()),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              self.tableLiq.item(i, 1).text(),
                                              self.tableLiq.item(i, 2).text(),
                                              "",
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                            elif (self.tableLiq.item(i, 1).text() == "" and self.tableLiq.item(i, 3).text() == "") or (
                                     self.tableLiq.item(i, 1).text() == "nan" and self.tableLiq.item(i, 3).text() == "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             re.sub(r'[\{\}]', "", self.tableLiq.item(i, 2).text()),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              "",
                                              self.tableLiq.item(i, 2).text(),
                                              "",
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                        elif self.tableLiq.item(i, 1) == None and self.tableLiq.item(i, 3) != None:
                            if self.tableLiq.item(i, 3).text() != "" or self.tableLiq.item(i, 3).text() != "nan":
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             re.sub(r'[\{\}]', "",self.tableLiq.item(i,2).text()) + "-" + self.tableLiq.item(i, 3).text(),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.',',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              "",
                                              self.tableLiq.item(i, 2).text(),
                                              self.tableLiq.item(i, 3).text(),
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                        elif self.tableLiq.item(i, 1) != None and self.tableLiq.item(i, 3) == None:
                            if self.tableLiq.item(i, 1).text() != ""  or self.tableLiq.item(i, 1).text() != "nan":
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             self.tableLiq.item(i, 1).text() + "-" + re.sub(r'[\{\}]', "",self.tableLiq.item(i,2).text()),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             str(float(self.tableLiq.item(i,5).text())).replace('.',','),
                                             str(round(temp9, 2)).replace('.',',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              self.tableLiq.item(i, 1).text(),
                                              self.tableLiq.item(i, 2).text(),
                                              "",
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                        elif self.tableLiq.item(i, 1) == None and self.tableLiq.item(i, 3) == None:
                            df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                         re.sub(r'[\{\}]', "", self.tableLiq.item(i,2).text()),
                                         sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                         round(temp7),
                                         self.tableLiq.item(i, 4).text(),
                                         str(float(round(temp6, 1))).replace('.', ','),
                                         str(float(round(temp8, 1))).replace('.', ','),
                                         self.tableLiq.item(i, 5).text(),
                                         str(round(temp9, 2)).replace('.',',')]
                            df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                          "",
                                          self.tableLiq.item(i, 2).text(),
                                          "",
                                          self.tableLiq.item(i, 4).text(),
                                          self.tableLiq.item(i, 5).text(),
                                          self.comboBoxCountLiq[i].currentText()]
                    if self.comboBoxCountLiq[i].currentText() == "RNA":
                        temp6 = (float(self.tableLiq.item(i,4).text().replace(',','.'))*(100/(((temp5['A']+temp5['a']) * 1.53)
                                                                                              + ((temp5['G'] + temp5['g']) * 1.18)
                                                                                              + ((temp5['C'] + temp5['c']) * 0.74)
                                                                                              + ((temp5['T'] + temp5['U'] + temp5['t'] + temp5['u']) * 0.93))) * 1000 / int(self.tableLiq.item(i,5).text()))
                        temp7 = (((temp5['A']+temp5['a'])*329.21)
                                 + ((temp5['T']+temp5['t'])*320.2)
                                 + ((temp5['C']+temp5['c'])*305.18)
                                 + ((temp5['G']+temp5['g'])*345.21)
                                 + ((temp5['U']+temp5['u'])*306.17)+159.0)
                        temp8 = (float(self.tableLiq.item(i,4).text().replace(',','.'))/temp6*1000)
                        temp9 = (float(self.tableLiq.item(i,5).text())*temp7*0.000001)
                        if (self.tableLiq.item(i, 1) != None and self.tableLiq.item(i, 3) != None):
                            if (self.tableLiq.item(i, 1).text() != "" and self.tableLiq.item(i, 3).text() != "") and (
                                    self.tableLiq.item(i, 1).text() != "nan" and self.tableLiq.item(i,
                                                                                                    3).text() != "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             self.tableLiq.item(i, 1).text() + "-" + re.sub(r'[\{\}]', "",
                                                                                            self.tableLiq.item(i,
                                                                                                               2).text()) + "-" + self.tableLiq.item(
                                                 i, 3).text(),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              self.tableLiq.item(i, 1).text(),
                                              self.tableLiq.item(i, 2).text(),
                                              self.tableLiq.item(i, 3).text(),
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                            elif (self.tableLiq.item(i, 1).text() == "" and self.tableLiq.item(i, 3).text() != "") or (
                                    self.tableLiq.item(i, 1).text() == "nan" and self.tableLiq.item(i,
                                                                                                    3).text() != "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             re.sub(r'[\{\}]', "",
                                                    self.tableLiq.item(i, 2).text()) + "-" + self.tableLiq.item(i,
                                                                                                                3).text(),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              "",
                                              self.tableLiq.item(i, 2).text(),
                                              self.tableLiq.item(i, 3).text(),
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                            elif (self.tableLiq.item(i, 1).text() != "" and self.tableLiq.item(i, 3).text() == "") or (
                                    self.tableLiq.item(i, 1).text() != "nan" and self.tableLiq.item(i,
                                                                                                    3).text() == "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             self.tableLiq.item(i, 1).text() + "-" + re.sub(r'[\{\}]', "",
                                                                                            self.tableLiq.item(i,
                                                                                                               2).text()),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              self.tableLiq.item(i, 1).text(),
                                              self.tableLiq.item(i, 2).text(),
                                              "",
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                            elif (self.tableLiq.item(i, 1).text() == "" and self.tableLiq.item(i, 3).text() == "") or (
                                    self.tableLiq.item(i, 1).text() == "nan" and self.tableLiq.item(i,
                                                                                                    3).text() == "nan"):
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             re.sub(r'[\{\}]', "", self.tableLiq.item(i, 2).text()),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              "",
                                              self.tableLiq.item(i, 2).text(),
                                              "",
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                        elif self.tableLiq.item(i, 1) == None and self.tableLiq.item(i, 3) != None:
                            if self.tableLiq.item(i, 3).text() != "" or self.tableLiq.item(i, 3).text() != "nan":
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             re.sub(r'[\{\}]', "",
                                                    self.tableLiq.item(i, 2).text()) + "-" + self.tableLiq.item(i,
                                                                                                                3).text(),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              "",
                                              self.tableLiq.item(i, 2).text(),
                                              self.tableLiq.item(i, 3).text(),
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                        elif self.tableLiq.item(i, 1) != None and self.tableLiq.item(i, 3) == None:
                            if self.tableLiq.item(i, 1).text() != "" or self.tableLiq.item(i, 1).text() != "nan":
                                df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                             self.tableLiq.item(i, 1).text() + "-" + re.sub(r'[\{\}]', "",
                                                                                            self.tableLiq.item(i,
                                                                                                               2).text()),
                                             sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                             round(temp7),
                                             self.tableLiq.item(i, 4).text(),
                                             str(float(round(temp6, 1))).replace('.', ','),
                                             str(float(round(temp8, 1))).replace('.', ','),
                                             self.tableLiq.item(i, 5).text(),
                                             str(round(temp9, 2)).replace('.', ',')]
                                df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                              self.tableLiq.item(i, 1).text(),
                                              self.tableLiq.item(i, 2).text(),
                                              "",
                                              self.tableLiq.item(i, 4).text(),
                                              self.tableLiq.item(i, 5).text(),
                                              self.comboBoxCountLiq[i].currentText()]
                        elif self.tableLiq.item(i, 1) == None and self.tableLiq.item(i, 3) == None:
                            df.loc[i] = [self.tableLiq.item(i, 0).text(),
                                         re.sub(r'[\{\}]', "", self.tableLiq.item(i, 2).text()),
                                         sum(map(str.isalpha, self.tableLiq.item(i, 2).text())),
                                         round(temp7),
                                         self.tableLiq.item(i, 4).text(),
                                         str(float(round(temp6, 1))).replace('.', ','),
                                         str(float(round(temp8, 1))).replace('.', ','),
                                         self.tableLiq.item(i, 5).text(),
                                         str(round(temp9, 2)).replace('.', ',')]
                            df2.loc[i] = [self.tableLiq.item(i, 0).text(),
                                          "",
                                          self.tableLiq.item(i, 2).text(),
                                          "",
                                          self.tableLiq.item(i, 4).text(),
                                          self.tableLiq.item(i, 5).text(),
                                          self.comboBoxCountLiq[i].currentText()]
                with pd.ExcelWriter(filename[0],engine='openpyxl') as writer:
                    df.to_excel(writer,sheet_name='Sheet1')
                    df2.to_excel(writer,sheet_name='Sheet2')
                if word_or_not == True:
                    doc = docx.Document('SampleLiquid.docx')
                    doc.tables  # a list of all tables in document
                    table = doc.tables[0]
                    # cells = table.rows[1].cells
                    for i in range(df.shape[0]):
                        table.add_row().cells

                        table.cell(i + 2, 0).text = str(i + 1)
                        table.cell(i + 2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 1).text = str(df.values[i, 0])
                        table.cell(i + 2, 1).paragraphs[0].runs[0].font.bold = True
                        table.cell(i + 2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        if (df2.values[i, 1] == "nan" and df2.values[i, 3] == "nan") or (df2.values[i, 1] == "" and df2.values[i, 3] == "") or (df2.values[i, 1] == None and df2.values[i, 3] == None):
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif (df2.values[i, 1] == "nan" and df2.values[i, 3] != "nan") or (df2.values[i, 1] == "" and df2.values[i, 3] != "") or (df2.values[i, 1] == None and df2.values[i, 3] != None):
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 3])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif (df2.values[i, 1] != "nan" and df2.values[i, 3] == "nan") or (df2.values[i, 1] != "" and df2.values[i, 3] == "") or (df2.values[i, 1] != None and df2.values[i, 3] == None):
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 1])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 1])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 2]))
                            table.cell(i + 2, 2).paragraphs[0].add_run("-")
                            table.cell(i + 2, 2).paragraphs[0].add_run(str(df2.values[i, 3])).font.bold = True
                            table.cell(i + 2, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                        table.cell(i + 2, 3).text = str(df.values[i, 2])
                        table.cell(i + 2, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 4).text = str(df.values[i, 3])
                        table.cell(i + 2, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 5).text = str(df.values[i, 4])
                        table.cell(i + 2, 5).paragraphs[0].runs[0].font.bold = True
                        table.cell(i + 2, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 6).text = str(df.values[i, 5])
                        table.cell(i + 2, 6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 7).text = str(df.values[i, 6])
                        table.cell(i + 2, 7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 8).text = str(df.values[i, 7])
                        table.cell(i + 2, 8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 9).text = str(df.values[i, 8])
                        table.cell(i + 2, 9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        table.cell(i + 2, 10).text = "ПААГ"
                        table.cell(i + 2, 10).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.style = 'Table Grid'
                    print('doc')
                    f_splited = filename[0].split('/')
                    fname = f_splited[len(f_splited) - 1][:-5]
                    print(fname)
                    doc.save(filename[0][:-5]+".docx")

                    doc = docx.Document('Stickers.docx')
                    table = doc.tables[0]
                    row = 0
                    column = 0
                    temp_row = 1
                    now = datetime.datetime.now()
                    for i in range(df.shape[0]):
                        table.cell(row, column).add_paragraph(
                            str(df.values[i, 0])).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        table.cell(row, column).paragraphs[1].runs[0].font.bold = True
                        if (df2.values[i, 1] == "nan" and df2.values[i, 3] == "nan") or (df2.values[i, 1] == "" and df2.values[i, 3] == "") or (df2.values[i, 1] == None and df2.values[i, 3] == None):
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(" 5'-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run("-3' ").font.bold = True
                        elif (df2.values[i, 1] == "nan" and df2.values[i, 3] != "nan") or (df2.values[i, 1] == "" and df2.values[i, 3] != "") or (df2.values[i, 1] == None and df2.values[i, 3] != None):
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(" 5'-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run(
                                "-" + str(df2.values[i, 3]) + " ").font.bold = True
                        elif (df2.values[i, 1] != "nan" and df2.values[i, 3] == "nan") or (
                                df2.values[i, 1] != "" and df2.values[i, 3] == "") or (
                                df2.values[i, 1] != None and df2.values[i, 3] == None):
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(
                                str(" " + df2.values[i, 1]) + "-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run("-3' ").font.bold = True
                        else:
                            table.cell(row, column).add_paragraph()
                            table.cell(row, column).paragraphs[2].add_run(
                                " " + str(df2.values[i, 1]) + "-").font.bold = True
                            table.cell(row, column).paragraphs[2].add_run(str(df2.values[i, 2]))
                            table.cell(row, column).paragraphs[2].add_run(
                                "-" + str(df2.values[i, 3]) + " ").font.bold = True
                        table.cell(row, column).add_paragraph(
                            str(df.values[i, 7] + " pmol/mkl")).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        table.cell(row, column).paragraphs[3].runs[0].font.bold = True
                        table.cell(row, column).add_paragraph(
                            " Stor.-20°C        " + now.strftime('%d %b %Y ')).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        column += 2
                        if column > 6:
                            row += 1
                            temp_row += 1
                            column = 0
                    for i in range(temp_row):
                        for j in range(0, 7, 2):
                            paragraphs = table.cell(i, j).paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size = Pt(6)
                    doc.save(filename[0][:-5] + 'Наклейки.docx')
        except:
            self._creatWarningWindow("Вы не заполнили все ячейки в таблице или не закрыли сохраняемый exel файл, данные не буддут корректно сохранены!")
    def saveExistFile(self):
        if self.tempFileLoc[0] != '':
            self.saveFile(self.tempFileLoc, False)
        else:
            self._creatWarningWindow("Файл не выбран!")


    def saveAsOnlyExel(self):
        filename = QFileDialog.getSaveFileName(self,"Save file","","Xlsx Files (*.xlsx)")
        if filename[0] != '':
            self.saveFile(filename, False)
            self.tempFileLoc = filename
            self.pathToFile(filename)
        print("Exel")
    def saveAsWithWord(self):
        filename = QFileDialog.getSaveFileName(self, "Save file", "", "Xlsx Files (*.xlsx)")
        if filename[0] != '':
            self.saveFile(filename, True)
            self.tempFileLoc = filename
            self.pathToFile(filename)
        print("Word")
        print(filename[0])
        print(filename[1])


    def pathToFile(self,path):
        tempFileLocForScreen = path[0].split("/")
        if len(tempFileLocForScreen) >= 7:
            self.lbl.setText("Путь к файлу: " + "..." + '/'.join(tempFileLocForScreen[-6:]))
            self.lbl.adjustSize()
        if len(tempFileLocForScreen) < 7:
            self.lbl.setText("Путь к файлу: " + '/'.join(tempFileLocForScreen[-6:]))
            self.lbl.adjustSize()
    def _connectAction(self):
        self.newAction.triggered.connect(self.newFile)
        self.sampleDry.triggered.connect(self.creatSampleDry)
        self.sampleLiq.triggered.connect(self.creatSampleLiquid)
        self.openAction.triggered.connect(self.openFile)
        self.saveAction.triggered.connect(self.saveExistFile)
        self.saveAsActionOnlyExel.triggered.connect(self.saveAsOnlyExel)
        self.saveAsActionWithWord.triggered.connect(self.saveAsWithWord)
        self.exitAction.triggered.connect(sys.exit)
    def _creatComboBox(self):
        self.comboBox = QComboBox(self)
        self.comboBox.addItems(["DNA", "RNA"])
        self.comboBox.setGeometry(472, 51, 48, 48)
        self.comboBox.textActivated.connect(self._creatDialogWindow)
    def _creatPushButtom2(self):
        self.pushButton2 = QPushButton(self)
        self.pushButton2.clicked.connect(self._feature1)
        self.pushButton2.setGeometry(20, 50, 50, 50)
        self.pushButton2.setText("+")
    def _creatPushButtom3(self):
        self.pushButton3 = QPushButton(self)
        self.pushButton3.clicked.connect(self._feature2)
        self.pushButton3.setGeometry(70, 50, 50, 50)
        self.pushButton3.setText("-")
    def _creatPushButtom4(self):
        self.pushButton4 = QPushButton(self)
        self.pushButton4.clicked.connect(self._feature3)
        self.pushButton4.setGeometry(520, 50, 50, 50)
        self.pushButton4.setIcon(QIcon("iconLiquid.png"))
    def _creatPushButtom5(self):
        self.pushButton5 = QPushButton(self)
        self.pushButton5.clicked.connect(self._feature4)
        self.pushButton5.setGeometry(570, 50, 50, 50)
        self.pushButton5.setIcon(QIcon("iconDry.png"))
    def _creatPushButtom6(self):
        self.pushButton5 = QPushButton(self)
        self.pushButton5.clicked.connect(self._creatInputWindow)
        self.pushButton5.setGeometry(422, 50, 50, 50)
        self.pushButton5.setText("OE")
    def _creatPushButtom7(self):
        self.pushButton6 = QPushButton(self)
        self.pushButton6.clicked.connect(self._creatInputWindow2)
        self.pushButton6.setGeometry(372, 50, 50, 50)
        self.pushButton6.setText("pml/mk")
        self.pushButton6.hide()

    def _feature1(self):
        if self.table.isVisible() == True:
            self.RowCount.append(self.RowCount[-1]+1)
            self.table.setRowCount(self.RowCount[-1])
            temp = QComboBox(self)
            self.comboBoxCount.append(temp)
            self.comboBoxCount[-1].addItems(["DNA","RNA"])
            self.table.setCellWidget(self.RowCount[-1]-1,5,self.comboBoxCount[-1])
        elif self.tableLiq.isVisible() == True:
            self.RowCountLiq.append(self.RowCountLiq[-1]+1)
            self.tableLiq.setRowCount(self.RowCountLiq[-1])
            temp = QComboBox(self)
            self.comboBoxCountLiq.append(temp)
            self.comboBoxCountLiq[-1].addItems(["DNA","RNA"])
            self.tableLiq.setCellWidget(self.RowCountLiq[-1]-1,6,self.comboBoxCountLiq[-1])

    def _feature2(self):
        if self.RowCount[-1] > 0 and self.table.isVisible() == True:
            self.RowCount.remove(self.RowCount[-1])
            self.table.setRowCount(self.RowCount[-1])
            self.comboBoxCount.remove(self.comboBoxCount[-1])
        elif self.RowCountLiq[-1] > 0 and self.tableLiq.isVisible() == True:
            self.RowCountLiq.remove(self.RowCountLiq[-1])
            self.tableLiq.setRowCount(self.RowCountLiq[-1])
            self.comboBoxCountLiq.remove(self.comboBoxCountLiq[-1])
    def _feature3(self):
        self.table.hide()
        self.tableLiq.show()
        self.pushButton6.show()
    def _feature4(self):
        self.pushButton6.hide()
        self.tableLiq.hide()
        self.table.show()
    def _creatTableWidgetDry(self):
        self.table = QTableWidget(self)
        self.table.setColumnCount(6)
        self.table.setHorizontalHeaderLabels(["Название","Зонд", "Последовательность","Зонд","OE", "DNA/RNA"])
        self.table.resizeColumnsToContents()
        self.table.setGeometry(20,110,600,300)
        self.comboBoxCount = []
    def _creatTableWidgetLiq(self):
        self.tableLiq = QTableWidget(self)
        self.tableLiq.setColumnCount(7)
        self.tableLiq.setHorizontalHeaderLabels(["Название","Зонд","Последовательность","Зонд","OE","помоль/мкг", "DNA/RNA"])
        self.tableLiq.resizeColumnsToContents()
        self.tableLiq.setGeometry(20,110,600,300)
        self.comboBoxCountLiq = []
        self.tableLiq.hide()
    def _creatDialogWindow(self):
        InDiaWin = QMessageBox()
        InDiaWin.setWindowTitle("Вопрос")
        InDiaWin.setText("Вы дейстивтельно хотите изменить все на " + self.comboBox.currentText() + "?")
        InDiaWin.setWindowIcon(QIcon("icon1.png"))
        InDiaWin.setIcon(QMessageBox.Question)
        InDiaWin.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
        InDiaWin.buttonClicked.connect(self._popup_button)
        InDiaWin.exec_()
    def _creatWarningWindow(self,text):
        InDiaWin = QMessageBox()
        InDiaWin.setWindowTitle("Внимание!")
        InDiaWin.setText(text)
        InDiaWin.setWindowIcon(QIcon("icon2.png"))
        InDiaWin.setIcon(QMessageBox.Warning)
        InDiaWin.setStandardButtons(QMessageBox.Ok)
        InDiaWin.exec_()
    def _creatInputWindow(self):
        text, ok = QInputDialog.getText(self,"Изменение ОЕ", "Введите ОЕ")
        if ok:
            if self.table.isVisible() == True:
                for i in range(0,self.RowCount[-1]):
                    self.table.setItem(i,4,QTableWidgetItem(text))
            elif self.tableLiq.isVisible() == True:
                for i in range(0,self.RowCountLiq[-1]):
                    self.tableLiq.setItem(i,4,QTableWidgetItem(text))
    def _creatInputWindow2(self):
        text, ok = QInputDialog.getText(self,"Изменение pmole/mkl", "Введите pmole/mkl")
        if ok:
                for i in range(0,self.RowCountLiq[-1]):
                    self.tableLiq.setItem(i,5,QTableWidgetItem(text))

    def _popup_button(self, btn):
            if btn.text() == "OK":
                if self.RowCount[-1] > 0 and self.table.isVisible() == True:
                    for i in range(0,self.RowCount[-1]):
                        self.comboBoxCount[i].setCurrentIndex(self.comboBox.currentIndex())
                elif self.RowCountLiq[-1] > 0 and self.tableLiq.isVisible() == True:
                    for i in range(0,self.RowCountLiq[-1]):
                        self.comboBoxCountLiq[i].setCurrentIndex(self.comboBox.currentIndex())
            else:
                pass
if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = Window()
    win.show()
    sys.exit(app.exec_())
