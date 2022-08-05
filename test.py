import docx
import math
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = docx.Document('Sample.docx')
table = doc.tables[0]
for i in range(2):
    table.add_row().cells

    table.cell(i+2,0).text = "хуй"
    table.cell(i + 2, 2).paragraphs[0].add_run("хуй").font.bold = True
    table.cell(i+2,2).paragraphs[0].add_run("залупа")
    table.cell(i + 2, 2).paragraphs[0].add_run("хуй").font.bold = True
    table.cell(i+2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
table.style = 'Table Grid'
#doc.save('test.docx')
df2 = pd.DataFrame(columns=['Название','Зонд','Последовательность','Зонд','OE','DNA/RNA'])
df2.loc[0] = [1,"CCTTT","AAAAAAAAAABBBBBBBBBBCCCCCCCCCC","G",1,None]
print(df2)
print(pd.isna(df2.values[0,5]))
if len(df2.values[0,2]) + len(df2.values[0,1]) + len(df2.values[0,3]) > 15 :
    num = df2.values[0,1] + df2.values[0,2] + df2.values[0,3]
    chunk_count = math.ceil(len(num)/15)
    n = 15
    print(n)
    chunks = [num[i:i + n] for i in range(0, len(num), n)]
    print(num)
    print(chunks[0])
    while True:
        tempy = df2.values[0,2]
        tempy2 = 15 - len(df2.values[0,1])
        print(tempy2)
        print(df2.values[0,1] + tempy[0:tempy2])
        print(tempy[tempy2:(tempy2 + 15)])
        break